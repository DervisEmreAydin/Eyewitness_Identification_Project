import docx

import data

doc = docx.Document('finalReportTemplate.docx')

list_police_data = list(data.dict_police_parameters.values())

# Populate Table 1
cellIndex = 0
policeDataIndex = 0
for index in range(6):
    doc.tables[0].cell(cellIndex, 1).text = list_police_data[policeDataIndex]
    cellIndex += 1
    policeDataIndex += 1

# Populate Table 2
cellIndex = 0
for index in range(3):
    doc.tables[1].cell(cellIndex, 1).text = list_police_data[policeDataIndex]
    cellIndex += 1
    policeDataIndex += 1

# Populate Table 3
cellIndex = 0
for index in range(3):
    doc.tables[2].cell(cellIndex, 1).text = list_police_data[policeDataIndex]
    cellIndex += 1
    policeDataIndex += 1

# Populate Table 4
cellIndex = 0
doc.tables[3].cell(0, 1).text = list_police_data[0]

# XXX Missing DATA
# Populate Table 5
cellIndex = 0
list_witness_data = list(data.dict_witness_parameters.values())
witnessDataIndex = 0
for cellIndex in range(3):
    doc.tables[4].cell(cellIndex, 1).text = list_witness_data[witnessDataIndex]
    cellIndex += 1
    witnessDataIndex += 1

### .... ###

### Populate Lineup Table 8
list_lineup = data.EyewitnessLineupList
rowIndex = 0
lineupIndex = 0
for rowIndex in range(2):
    for columnIndex in range(3):
        paragraph = doc.tables[7].cell(rowIndex, columnIndex).paragraphs[0]
        run = paragraph.add_run()
        lineupIndex = (rowIndex * 3) + columnIndex
        run.add_picture(list_lineup[lineupIndex], width=1400000, height=1400000)

# Populate Table 9 with the Image identified by the witness
if data.witness_image_choice == "-":
    doc.tables[8].cell(0, 0).text = "Hicbiri uymadi"
else:
    paragraph = doc.tables[8].cell(0, 0).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(data.witness_image_choice, width=1400000, height=1400000)

# Populate table 10 with confidence rate
doc.tables[9].cell(0, 1).text = data.witness_confidence
doc.save('finalReport.docx')
